import pandas as pd
from docx import Document
import os
import re
from docx.shared import Pt

TEMPLATE_FILE = "template_BBNT_GOI_CAO_CAP.docx"
EXCEL_FILE = "BBNT_DATA_TEMPLATE.xlsx"
OUTPUT_FOLDER = "OUTPUT_BBNT"

os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def clean_name(text):
    return re.sub(r'[\\/:*?"<>|]', '', str(text))

df = pd.read_excel(EXCEL_FILE)

# # def replace_all(doc, data):
#     # Paragraphs
#     for p in doc.paragraphs:
#         for run in p.runs:
#             for k, v in data.items():
#                 if k in run.text:
#                     run.text = run.text.replace(k, str(v))

#     # Tables
#     for table in doc.tables:
#         for row in table.rows:
#             for cell in row.cells:
#                 for p in cell.paragraphs:
#                     for run in p.runs:
#                         for k, v in data.items():
#                             if k in run.text:
#                                 run.text = run.text.replace(k, str(v))
def replace_all(doc, data):
    def replace_in_paragraph(p):
        if not p.runs:
            return

        full_text = "".join(run.text for run in p.runs)

        replaced = False
        for k, v in data.items():
            if k in full_text:
                full_text = full_text.replace(k, str(v))
                replaced = True

        if not replaced:
            return

        # Ghi lại text vào các run cũ (giữ format)
        idx = 0
        for run in p.runs:
            run_len = len(run.text)
            run.text = full_text[idx: idx + run_len]
            idx += run_len

        # Nếu còn dư text thì append vào run cuối
        if idx < len(full_text):
            p.runs[-1].text += full_text[idx:]

    # Paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

for _, r in df.iterrows():
    doc = Document(TEMPLATE_FILE)
    so_hd = str(r["SO_HOP_DONG"]).zfill(4)

    data = {
        "{{SO_HOP_DONG}}": so_hd,
        "{{NGAY_KY_HOP_DONG}}": r["NGAY_KY_HOP_DONG"],
        "{{NGAY_LAP_BB}}": r["NGAY_LAP_BB"],
        "{{TEN_BEN_B}}": r["TEN_BEN_B"],
        "{{DAI_DIEN}}": r["DAI_DIEN"],
        "{{DIA_CHI_CCCD}}": r["DIA_CHI_CCCD"],
        "{{MST_CCCD}}": r["MST_CCCD"],
        "{{CHUC_VU}}": r["CHUC_VU"],
        "{{DIA_CHI_SAN}}": r["DIA_CHI_SAN"],
        "{{DIEN_THOAI}}": r["DIEN_THOAI"],
        "{{EMAIL}}": r["EMAIL"],
        "{{THOI_HAN}}": r["THOI_HAN"],
    }

    replace_all(doc, data)

    file_name = (
        f"{so_hd}"
        f"_BBBGNT_{clean_name(r['TEN_SAN'])}"
        f"_GOI_CAO_CAP.docx"
    )

    doc.save(os.path.join(OUTPUT_FOLDER, file_name))

print("✅ Chào mừng bạn đã hoàn thành việc tạo biên bản nghiệm thu gói cao cấp tự động của Minh Toàn!")
